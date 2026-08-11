"""
Generate DOCX proposal for IsyaratLive — GEMASTIK XIX 2026
Clean formatting, no emojis, 50% progress milestone scope, updated screenshots.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os

# === Screenshot paths (Updated clean screenshots without camera feed/emojis) ===
SCREENSHOTS_DIR = r"C:\Users\Habib\.gemini\antigravity-ide\brain\29f11f76-7bbf-4c3f-9cf5-90d4e5dcf59d"
LANDING_IMG  = os.path.join(SCREENSHOTS_DIR, "screenshot_landing_1786436168740.png")
MODE2_IMG    = os.path.join(SCREENSHOTS_DIR, "screenshot_room_1786436198018.png")
DICT_IMG     = os.path.join(SCREENSHOTS_DIR, "screenshot_dictionary_1786436212365.png")
ACCURACY_IMG = os.path.join(SCREENSHOTS_DIR, "screenshot_accuracy_1786436237326.png")

OUTPUT_PATH_1 = r"D:\Penyimpanan Utama\Documents\python\IsyaratLive\docs\IsyaratLive_Proposal_GEMASTIK_XIX.docx"
OUTPUT_PATH_2 = r"D:\Penyimpanan Utama\Documents\python\IsyaratLive\docs\Proposal_IsyaratLive_GEMASTIK_XIX.docx"

# === Helpers ===
def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None, header_color="1B3A4B"):
    """Add a formatted table with colored header."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)
    
    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F0F4F8")
    
    # Set column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    
    return table

def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x4B)
    return h

def para(doc, text, bold=False, size=11, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(doc, text, level=0, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(size)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * level)
    return p

def add_image_with_caption(doc, img_path, caption, width=Inches(5.8)):
    """Add image centered with a caption below."""
    if not os.path.exists(img_path):
        para(doc, f"[Screenshot tidak ditemukan: {img_path}]", size=9)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=width)
    
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.size = Pt(9)
    r.italic = True
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    cap.paragraph_format.space_after = Pt(12)

# ============================================================
# BUILD DOCUMENT
# ============================================================
doc = Document()

# -- Page setup --
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ============================================================
# COVER PAGE
# ============================================================
for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PROPOSAL PENGEMBANGAN PERANGKAT LUNAK\n(TAHAP KEMAJUAN 50%)")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x4B)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("IsyaratLive")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x0E, 0x7C, 0x7B)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Penerjemah Bahasa Isyarat Indonesia (BISINDO)\nReal-Time Berbasis AI (CV + LLM)")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("GEMASTIK XIX 2026\nDivisi VIII: Pengembangan Perangkat Lunak")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x4B)

for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("2026")
run.font.size = Pt(12)
run.bold = True

doc.add_page_break()

# ============================================================
# DAFTAR ISI
# ============================================================
heading(doc, "DAFTAR ISI", level=1)

toc_items = [
    ("a)", "Judul/Nama Perangkat Lunak", "1"),
    ("b)", "Latar Belakang Ide Perangkat Lunak", "1"),
    ("c)", "Tujuan dan Manfaat Dikembangkannya Perangkat Lunak", "4"),
    ("d)", "Batasan Perangkat Lunak yang Dikembangkan", "6"),
    ("e)", "Metodologi Pengembangan Perangkat Lunak", "7"),
    ("f)", "Analisis Kebutuhan dan Desain Solusi Perangkat Lunak", "10"),
    ("g)", "Implementasi Perangkat Lunak", "14"),
    ("h)", "Screenshot Mockup Interface Perangkat Lunak", "22"),
    ("i)", "Dokumentasi Cara Penggunaan Perangkat Lunak", "25"),
    ("", "Kelebihan Dibandingkan Perangkat Lunak Serupa", "28"),
    ("", "Daftar Komponen dan Lisensi", "29"),
    ("", "Referensi dan Sitasi", "30"),
]

for num, title, page in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f"{num} {title}" if num else title)
    run.font.size = Pt(11)
    run2 = p.add_run(f"\t{page}")
    run2.font.size = Pt(11)

doc.add_page_break()

# ============================================================
# a) JUDUL
# ============================================================
heading(doc, "a) Judul/Nama Perangkat Lunak", level=1)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("IsyaratLive")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x0E, 0x7C, 0x7B)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Penerjemah Bahasa Isyarat Indonesia (BISINDO) Real-Time Berbasis AI (CV + LLM)")
run.font.size = Pt(12)
run.italic = True

doc.add_paragraph()

# ============================================================
# b) LATAR BELAKANG
# ============================================================
heading(doc, "b) Latar Belakang Ide Perangkat Lunak", level=1)

heading(doc, "Urgensi Permasalahan", level=2)
para(doc, "Indonesia merupakan salah satu negara dengan jumlah penyandang disabilitas pendengaran (Tuli) tertinggi di Asia Tenggara. Berdasarkan data Badan Pusat Statistik (BPS), terdapat jutaan penyandang disabilitas pendengaran dan wicara di Indonesia yang membutuhkan sarana komunikasi alternatif dalam kehidupan sehari-hari. Ironisnya, jumlah juru bahasa isyarat bersertifikat di Indonesia sangat terbatas dan tidak selalu tersedia di titik-titik layanan publik seperti puskesmas, kantor kelurahan, kepolisian, sekolah inklusi, maupun lingkungan keluarga campuran (keluarga dengan anggota Tuli dan dengar).")

para(doc, "Ketiadaan juru bahasa isyarat di titik-titik layanan publik ini menciptakan hambatan komunikasi yang serius, menyebabkan penyandang Tuli sering kali tidak mendapatkan layanan yang setara, kesalahpahaman dalam proses administrasi, bahkan isolasi sosial di lingkungan keluarga sendiri.")

heading(doc, "Keterbatasan Solusi yang Ada", level=2)
para(doc, "Aplikasi penerjemah bahasa isyarat yang sudah ada di Indonesia saat ini memiliki beberapa keterbatasan fundamental:")

bullet(doc, "Fokus pada SIBI, bukan BISINDO. Sebagian besar solusi yang ada menerjemahkan Sistem Isyarat Bahasa Indonesia (SIBI) — sebuah sistem buatan yang mengikuti struktur tata bahasa Indonesia lisan. Padahal, komunitas Tuli Indonesia secara alami menggunakan BISINDO (Bahasa Isyarat Indonesia) yang memiliki tata bahasa sendiri dengan struktur topic-comment, berbeda dari bahasa Indonesia lisan.")
bullet(doc, "Terbatas pada alfabet/angka statis. Aplikasi yang mengklaim mendukung BISINDO umumnya hanya mampu mendeteksi isyarat statis (satu frame) berupa alfabet atau angka — bukan rangkaian kata dinamis yang digunakan dalam percakapan nyata.")
bullet(doc, "Tidak ada lapisan pemahaman bahasa. Solusi yang ada berhenti pada tahap \"deteksi kata/huruf\" tanpa kemampuan menyusun output menjadi kalimat yang dapat dipahami secara natural oleh orang dengar. Output berupa tumpukan kata acak seperti \"SAYA MAKAN SUDAH TADI\" sangat berbeda dari kalimat natural \"Saya sudah makan tadi.\"")

heading(doc, "Ide dan Solusi IsyaratLive", level=2)
para(doc, "IsyaratLive hadir sebagai solusi inovatif yang menggabungkan dua lapis kecerdasan buatan (AI) yang saling melengkapi dan sama-sama esensial:")

bullet(doc, "Lapisan pertama — Computer Vision (CV): Mendeteksi gerakan tangan pengguna secara real-time melalui kamera web, mengekstrak 21 titik landmark per tangan menggunakan MediaPipe Hand Landmarker, dan mengklasifikasikan rangkaian gerakan menjadi gloss (representasi kata dalam BISINDO) menggunakan model TensorFlow.js (Model v1/v2) yang dilatih pada dataset WL-BISINDO.")
bullet(doc, "Lapisan kedua — Large Language Model (LLM): Menyusun ulang rangkaian gloss yang terdeteksi menjadi kalimat Bahasa Indonesia natural yang dapat dipahami dengan mudah oleh orang dengar, mengakomodasi perbedaan struktur tata bahasa BISINDO (topic-comment) dan bahasa Indonesia lisan.")

para(doc, "Kedua lapisan ini bersifat esensial dan tidak tergantikan — tanpa lapisan CV, sistem tidak bisa membaca isyarat sama sekali; tanpa lapisan LLM, output hanya berupa tumpukan kata acak yang sulit dipahami. Inilah yang membedakan IsyaratLive dari seluruh solusi penerjemah BISINDO yang ada saat ini.")

para(doc, "IsyaratLive dirancang sebagai aplikasi web berbasis React 19, Vite 8, dan Tailwind CSS 4 dengan tampilan light mode yang bersih, modern, dan tanpa emoticon. Aplikasi mendukung dua skenario penggunaan: Room Lokal (tatap muka satu perangkat) dan Room Remote (panggilan video P2P WebRTC 1-lawan-1 antar dua lokasi).")

# ============================================================
# c) TUJUAN DAN MANFAAT
# ============================================================
doc.add_page_break()
heading(doc, "c) Tujuan dan Manfaat Dikembangkannya Perangkat Lunak", level=1)

heading(doc, "Tujuan", level=2)

bullet(doc, "Menerjemahkan gerakan BISINDO menjadi kalimat Bahasa Indonesia natural secara real-time. Bukan hanya mendeteksi alfabet atau kata lepas, melainkan rangkaian kata yang disusun menjadi kalimat bermakna menggunakan kecerdasan buatan (CV + LLM).")
bullet(doc, "Menerjemahkan ucapan/teks Bahasa Indonesia menjadi rangkaian visual isyarat BISINDO. Menyediakan komunikasi dua arah yang memungkinkan orang dengar menyampaikan pesan dalam bentuk yang dapat dipahami oleh penyandang Tuli.")
bullet(doc, "Membuktikan bahwa AI (Computer Vision + Large Language Model) adalah komponen inti yang tidak tergantikan dalam penerjemahan bahasa isyarat level-kata — bukan sekadar fitur tambahan atau tempelan pada aplikasi biasa.")
bullet(doc, "Menyediakan alat bantu komunikasi yang aksesibel dan murah — cukup membuka browser, tanpa memerlukan instalasi aplikasi khusus, perangkat keras tambahan, atau biaya berlangganan.")

heading(doc, "Manfaat", level=2)

add_styled_table(doc,
    ["Sasaran", "Manfaat"],
    [
        ["Penyandang Tuli", "Dapat berkomunikasi dengan orang dengar secara mandiri tanpa bergantung pada ketersediaan juru bahasa isyarat manusia, baik di rumah, sekolah, maupun layanan publik"],
        ["Orang dengar (petugas layanan publik, keluarga, guru)", "Dapat memahami pesan dari penyandang Tuli dan merespons dengan cara yang dapat dipahami, meningkatkan kualitas pelayanan dan hubungan interpersonal"],
        ["Institusi (sekolah inklusi, puskesmas, kantor kelurahan)", "Memiliki alat bantu komunikasi yang murah, tidak memerlukan instalasi khusus, dan dapat digunakan di perangkat apa pun yang memiliki browser dan kamera"],
        ["Komunitas Tuli Indonesia", "Meningkatkan visibilitas dan akseptasi BISINDO sebagai bahasa alami komunitas Tuli, dibandingkan SIBI yang merupakan konstruksi artifisial"],
        ["Pengembangan ilmu pengetahuan", "Menjadi bukti konsep (proof of concept) bahwa kombinasi CV + LLM dapat mengatasi gap linguistik antara bahasa isyarat dan bahasa lisan"],
    ],
    col_widths=[5, 12]
)

# ============================================================
# d) BATASAN
# ============================================================
doc.add_page_break()
heading(doc, "d) Batasan Perangkat Lunak yang Dikembangkan", level=1)

bullet(doc, "Kosakata terbatas pada 32 kata BISINDO. Versi kompetisi ini mendukung 32 kosakata yang mencakup kata tanya, keseharian, keluarga, warna, dan waktu. Kosakata ini cukup untuk menyusun kalimat sederhana yang bermakna (mis. \"Saya makan pagi\", \"Terima kasih\", \"Kapan teman datang?\").")
bullet(doc, "Varian regional BISINDO Banten. Model dilatih menggunakan dataset WL-BISINDO yang direkam oleh penanda tangan dari regional Banten. Variasi gerakan dari daerah lain mungkin tidak sepenuhnya dikenali dengan tingkat akurasi yang sama.")
bullet(doc, "Teks-ke-isyarat berbasis dictionary, bukan generatif. Mode terjemahan teks ke isyarat menggunakan 32 file video peragaan isyarat yang sudah direkam sebelumnya (dictionary-based), bukan avatar 3D animasi generatif.")
bullet(doc, "Kalimat sederhana (satu klausa). Pada versi MVP, penerjemahan difokuskan pada kalimat sederhana — belum mendukung kalimat kompleks multi-klausa.")
bullet(doc, "Membutuhkan kamera dan browser modern. Aplikasi memerlukan perangkat dengan kamera (laptop/smartphone) dan browser yang mendukung WebGL/WebGPU, MediaPipe Tasks-Vision, serta Web Speech API.")
bullet(doc, "Lapisan LLM membutuhkan koneksi internet. Penyusunan kalimat natural memerlukan koneksi ke server backend. Tanpa koneksi internet, sistem tetap berfungsi dalam mode degradasi (menampilkan gloss mentah secara lokal tanpa crash).")
bullet(doc, "Cakupan dokumen proposal (Kemajuan 50%). Proposal ini mencakup progres pengembangan 50%. Tampilan tangkapan layar (screenshot) fitur kamera real-time disajikan pada tahap integrasi produk final 100%, sementara dokumen ini menyajikan screenshot komponen operasional utama yang telah berfungsi.")

# ============================================================
# e) METODOLOGI
# ============================================================
doc.add_page_break()
heading(doc, "e) Metodologi Pengembangan Perangkat Lunak", level=1)

heading(doc, "Pendekatan: Iteratif-Inkremental dengan Decision Gate", level=2)
para(doc, "IsyaratLive dikembangkan menggunakan pendekatan iteratif-inkremental yang disesuaikan dengan karakteristik proyek berbasis AI — dimana keberhasilan setiap fase bergantung pada validasi hasil fase sebelumnya. Setiap fase memiliki decision gate (titik keputusan) yang harus dipenuhi sebelum melanjutkan ke fase berikutnya.")

heading(doc, "Timeline Pengembangan (±3 Bulan)", level=2)

add_styled_table(doc,
    ["Fase", "Periode", "Fokus", "Decision Gate"],
    [
        ["Fase 0", "Minggu 1", "Setup infrastruktur proyek (repo, scaffold frontend React+Vite, backend Express, environment ML Python)", "Seluruh npm install bersih, dev server berjalan"],
        ["Fase 1", "Bulan 1", "Dataset & Model — unduh WL-BISINDO, split Signer-Independent, ekstraksi landmark, fine-tune model (v1/v2), ekspor ke TensorFlow.js", "Model klasifikasi gloss berjalan real-time di browser (evaluasi sanity check 32 label)"],
        ["Fase 2", "Bulan 2, Minggu 1–2", "Backend & Integrasi LLM — endpoint /api/normalize, integrasi 9Router, setup MySQL untuk riwayat percakapan", "Endpoint /api/normalize mengembalikan kalimat natural dari input gloss"],
        ["Fase 3", "Bulan 2, Minggu 3–4", "Frontend & Room Lokal — kamera -> landmark -> klasifikasi -> normalisasi -> teks -> suara + shared chat feed", "Pipeline lengkap dan test suite Vitest (31 frontend test, 10 backend test) passing 100%"],
        ["Fase 4", "Bulan 3, Minggu 1–2", "Mode 2 & Room Remote WebRTC — dictionary 32 video, teks->isyarat, signaling Socket.io, redesign UI light mode", "Kedua mode (Lokal & Remote) berfungsi dan diuji"],
        ["Fase 5", "Bulan 3, Minggu 3–4", "Deployment & Finalisasi Demo — persiapkan stack Docker (Dockerfile, Caddy, docker-compose), deployment VPS HTTPS", "Sistem 100% siap di environment produksi"],
    ],
    col_widths=[2, 3.5, 6, 5.5]
)

doc.add_paragraph()
heading(doc, "Strategi Mitigasi Risiko", level=2)

add_styled_table(doc,
    ["Risiko", "Dampak", "Mitigasi"],
    [
        ["Akurasi model klasifikasi rendah setelah fine-tune", "Mode 1 tidak bisa didemokan secara meyakinkan", "Gunakan Model v1/v2 yang stabil dengan 126D landmark dan resampling sequence temporal"],
        ["Ekspor ke TensorFlow.js gagal/tidak kompatibel", "Inferensi tidak bisa di browser", "Gunakan format ekspor Sequential Keras yang teruji pada tfjs 4.22"],
        ["Latensi LLM saat demo", "Pengalaman pengguna terganggu", "9Router menyediakan fallback multi-provider otomatis; gunakan model llama-3.1-8b-instant"],
        ["Performa real-time di browser lambat", "UX terganggu", "Jalankan inferensi di WebGL/WebGPU delegate dan optimalkan interval sampling frame"],
        ["Koneksi internet terputus saat demo", "LLM tidak bisa dipanggil", "Mode degradasi anggun (offline): tampilkan gloss mentah secara lokal di browser tanpa crash"],
        ["Format data tidak konsisten pada dictionary", "Video kata tidak dapat diputar", "Validasi otomatis 32 file video .mp4 di frontend/public/dictionary/ sejajar dengan GLOSS_LABELS via unit test"],
    ],
    col_widths=[5, 4.5, 7.5]
)

# ============================================================
# f) ANALISIS KEBUTUHAN DAN DESAIN SOLUSI
# ============================================================
doc.add_page_break()
heading(doc, "f) Analisis Kebutuhan dan Desain Solusi Perangkat Lunak", level=1)

heading(doc, "Analisis Kebutuhan Fungsional", level=2)

add_styled_table(doc,
    ["ID", "Kebutuhan Fungsional", "Prioritas"],
    [
        ["FR-01", "Sistem dapat mendeteksi gerakan tangan pengguna melalui kamera web secara real-time", "Wajib (MVP)"],
        ["FR-02", "Sistem dapat mengekstrak 21 titik landmark per tangan menggunakan MediaPipe Hand Landmarker", "Wajib (MVP)"],
        ["FR-03", "Sistem dapat mengklasifikasikan rangkaian landmark menjadi gloss (kata BISINDO) menggunakan model AI", "Wajib (MVP)"],
        ["FR-04", "Sistem dapat menyusun rangkaian gloss menjadi kalimat Bahasa Indonesia natural menggunakan LLM", "Wajib (MVP)"],
        ["FR-05", "Sistem dapat membacakan kalimat hasil terjemahan menjadi suara (text-to-speech)", "Wajib (MVP)"],
        ["FR-06", "Sistem dapat menyimpan dan menampilkan riwayat percakapan dalam feed obrolan Room", "Wajib (MVP)"],
        ["FR-07", "Sistem tetap berfungsi (mode degradasi anggun) saat koneksi internet terputus", "Wajib (MVP)"],
        ["FR-08", "Sistem dapat menerjemahkan teks/suara menjadi rangkaian video isyarat BISINDO (Mode 2)", "Wajib (MVP)"],
        ["FR-09", "Sistem menyediakan kamus kosakata 32 label BISINDO interaktif yang dapat dijelajahi pengguna", "Wajib (MVP)"],
        ["FR-10", "Sistem menyediakan fasilitas panggilan video P2P WebRTC 1-lawan-1 (Room Remote)", "Opsional"],
        ["FR-11", "Sistem menyediakan panel Uji Akurasi Model untuk sanity check performa model TFJS di browser", "Opsional"],
    ],
    col_widths=[2, 11, 4]
)

doc.add_paragraph()
heading(doc, "Analisis Kebutuhan Non-Fungsional", level=2)

add_styled_table(doc,
    ["ID", "Kebutuhan Non-Fungsional", "Target"],
    [
        ["NFR-01", "Latensi end-to-end dari gerakan isyarat sampai suara keluar", "< 2–3 detik"],
        ["NFR-02", "Akurasi klasifikasi gloss pada 32 kata target (evaluasi sanity check 32 label)", "80–90%"],
        ["NFR-03", "Aplikasi dapat berjalan di browser modern (Chrome, Firefox, Edge)", "Cross-browser"],
        ["NFR-04", "Desain antarmuka bersih (light mode, tanpa emoticon, warna aksen teal terpusat)", "UI/UX Standar"],
        ["NFR-05", "Koneksi aman (HTTPS) wajib untuk akses kamera browser", "Let's Encrypt / Docker"],
        ["NFR-06", "Test coverage unit testing frontend dan backend", "31 FE tests, 10 BE tests passing"],
    ],
    col_widths=[2, 11, 4]
)

doc.add_paragraph()
heading(doc, "Desain Arsitektur Sistem", level=2)

para(doc, "IsyaratLive menggunakan arsitektur client-heavy dimana komputasi AI utama (deteksi landmark tangan dan klasifikasi gloss) berjalan langsung di browser pengguna menggunakan WebAssembly dan WebGL/WebGPU, sementara backend Node.js + Express menangani normalisasi bahasa (LLM 9Router), signaling WebRTC, dan penyimpanan riwayat MySQL.")

# Architecture Table
arch_table = doc.add_table(rows=2, cols=1)
arch_table.alignment = WD_TABLE_ALIGNMENT.CENTER
arch_table.style = 'Table Grid'

cell = arch_table.rows[0].cells[0]
cell.text = ""
set_cell_shading(cell, "E8F4FD")
p = cell.paragraphs[0]
r = p.add_run("BROWSER (Client — React 19 + Vite 8 + Tailwind CSS 4)\n")
r.bold = True
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x1B, 0x3A, 0x4B)
r = p.add_run(
    "• Kamera -> MediaPipe Hand Landmarker (JS/WASM, real-time, GPU delegate)\n"
    "• Buffer landmark sequence (30 frame / ~1 detik)\n"
    "• Model klasifikasi gloss (TensorFlow.js Model v1/v2)\n"
    "  -> Output: Array gloss [\"SAYA\", \"MAKAN\", \"SUDAH\", \"TADI\"]\n"
    "• Kirim gloss ke backend via HTTPS POST /api/normalize\n"
    "• Terima kalimat natural -> Tampilkan teks + Web Speech API (TTS/STT)\n"
    "• Room Remote WebRTC P2P Client + Kamus 32 Label + Panel Uji Akurasi"
)
r.font.size = Pt(9)

cell = arch_table.rows[1].cells[0]
cell.text = ""
set_cell_shading(cell, "FFF8E7")
p = cell.paragraphs[0]
r = p.add_run("BACKEND (VPS / Docker Container — Node.js + Express 5)\n")
r.bold = True
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x1B, 0x3A, 0x4B)
r = p.add_run(
    "• POST /api/normalize { gloss: [...] } -> 9Router (LLM gateway llama-3.1-8b-instant)\n"
    "  -> Return { text: \"kalimat natural\" }\n"
    "• MySQL — riwayat percakapan (/api/history)\n"
    "• Signaling Socket.io (/socket.io) untuk Room Remote WebRTC P2P\n"
    "• Caddy Edge Reverse Proxy / HTTPS Let's Encrypt"
)
r.font.size = Pt(9)

doc.add_paragraph()

# ============================================================
# g) IMPLEMENTASI
# ============================================================
doc.add_page_break()
heading(doc, "g) Implementasi Perangkat Lunak", level=1)

heading(doc, "Tech Stack", level=2)

add_styled_table(doc,
    ["Layer", "Teknologi", "Keterangan"],
    [
        ["Deteksi Landmark Tangan", "@mediapipe/tasks-vision (JS)", "Real-time, berjalan di browser, 21 titik landmark per tangan via WebGL/WebGPU"],
        ["Klasifikasi Gloss", "TensorFlow.js (Model v1/v2)", "Model sequence classifier berbasis landmark 126D, berjalan 100% di browser"],
        ["Normalisasi Bahasa", "9Router -> LLM (llama-3.1-8b-instant)", "Backend-only proxy, API key tersimpan aman di server"],
        ["Text-to-Speech & STT", "Web Speech API (built-in browser)", "Gratis, mendukung Bahasa Indonesia"],
        ["Frontend", "React 19 + Vite 8 + Tailwind CSS 4 + TypeScript", "SPA modern, responsif, light mode bersih tanpa emoticon"],
        ["Backend", "Node.js + Express 5 + TypeScript", "Proxy ke 9Router, endpoint normalisasi & riwayat"],
        ["Signaling WebRTC", "Socket.io 4.8", "Signaling peer-to-peer panggilan video Room Remote"],
        ["Testing", "Vitest + Supertest", "31 frontend unit test & 10 backend test passing 100%"],
        ["Database", "MySQL (mysql2)", "Riwayat percakapan"],
        ["Deployment", "Docker + Caddy + HTTPS Let's Encrypt", "Stack kontainer Docker siap deploy ke VPS"],
    ],
    col_widths=[4, 6, 7]
)

doc.add_paragraph()
heading(doc, "Dataset: WL-BISINDO (32 Kosakata)", level=2)

para(doc, "IsyaratLive menggunakan dataset publik WL-BISINDO (Word-Level BISINDO) sebagai sumber data pelatihan model klasifikasi gloss.")

add_styled_table(doc,
    ["Label", "Kata", "Label", "Kata", "Label", "Kata", "Label", "Kata"],
    [
        ["0", "Air",    "8",  "Motor",        "16", "Mengapa",    "24", "Datang"],
        ["1", "Belajar","9",  "Saya",         "17", "Bagaimana",  "25", "Teman"],
        ["2", "Cari",   "10", "Terima kasih", "18", "Merah",      "26", "Keluarga"],
        ["3", "Hari",   "11", "Tuli",         "19", "Kuning",     "27", "Rumah"],
        ["4", "Ingat",  "12", "Apa",          "20", "Hijau",      "28", "Pagi"],
        ["5", "Lagi",   "13", "Siapa",        "21", "Hitam",      "29", "Siang"],
        ["6", "Maaf",   "14", "Kapan",        "22", "Dengar",     "30", "Sore"],
        ["7", "Makan",  "15", "Di mana",      "23", "Berangkat",  "31", "Malam"],
    ],
    col_widths=[1.5, 3, 1.5, 3, 1.5, 3, 1.5, 3]
)

doc.add_paragraph()
heading(doc, "Struktur Proyek", level=2)

struct_text = (
    "isyaratlive/\n"
    "├── ml/                          # Training model (Python, offline)\n"
    "│   ├── dataset/                 # Dataset WL-BISINDO\n"
    "│   ├── preprocessing/           # Ekstraksi landmark (MediaPipe Python)\n"
    "│   ├── training/                # Fine-tuning model Siformer/SPOTER\n"
    "│   └── export/                  # Konversi model ke TensorFlow.js\n"
    "├── frontend/                    # React 19 + Vite 8 + TypeScript\n"
    "│   ├── src/components/          # CameraCapture, LandmarkDetector, GlossClassifier,\n"
    "│   │                            # ChatDisplay, DictionaryModal, AccuracyTestPanel\n"
    "│   ├── src/modes/               # SignToTextMode, TextToSignMode\n"
    "│   ├── src/rooms/               # RoomLocal.tsx, RoomRemote.tsx\n"
    "│   └── public/dictionary/       # 32 File video .mp4 peragaan isyarat\n"
    "├── backend/                     # Node.js + Express 5 + TypeScript\n"
    "│   ├── src/routes/              # normalize.ts, history.ts\n"
    "│   └── src/signaling.ts         # Socket.io WebRTC signaling\n"
    "├── deploy/                      # Caddyfile & Docker compose configuration\n"
    "└── docs/                        # Proposal & PRD documentation"
)
p = doc.add_paragraph()
run = p.add_run(struct_text)
run.font.name = 'Consolas'
run.font.size = Pt(8)

# ============================================================
# h) SCREENSHOT MOCKUP INTERFACE
# ============================================================
doc.add_page_break()
heading(doc, "h) Screenshot Mockup Interface Perangkat Lunak", level=1)

para(doc, "Berikut adalah screenshot antarmuka IsyaratLive yang diambil langsung dari aplikasi yang sedang berjalan (development server) dengan desain UI light mode bersih tanpa emoticon.")

# Catatan Fitur Kamera 50%
note_table = doc.add_table(rows=1, cols=1)
note_table.alignment = WD_TABLE_ALIGNMENT.CENTER
note_table.style = 'Table Grid'
cell_note = note_table.rows[0].cells[0]
set_cell_shading(cell_note, "E8F4FD")
p_note = cell_note.paragraphs[0]
r_note = p_note.add_run(
    "CATATAN MENGENAI FITUR KAMERA (MODE 1 — ISYARAT -> TEKS & SUARA):\n"
    "Pengembangan perangkat lunak ini saat ini berada pada tahap kemajuan minimal 50% (Tahap Proposal GEMASTIK XIX). "
    "Fitur terjemahan isyarat real-time melalui kamera (Mode 1) telah selesai dirancang pada level arsitektur (MediaPipe Hand Landmarker + Model TensorFlow.js v1/v2) dan sedang dalam proses finalisasi integrasi. "
    "Khusus untuk dokumen proposal kemajuan 50% ini, screenshot tampilan kamera belum dicantumkan secara visual. "
    "Fitur kamera tersebut akan dicantumkan dan didemokan secara penuh pada tahap produk final 100%. "
    "Dokumen ini menyajikan antarmuka komponen utama yang telah berfungsi operasional meliputi Beranda (Landing Page), Room Percakapan, Mode 2 (Teks -> Isyarat), Kamus 32 Label Kosakata, serta Panel Uji Akurasi Model."
)
r_note.font.size = Pt(9.5)
r_note.italic = True

doc.add_paragraph()
heading(doc, "1. Tampilan Landing Page (Beranda Utama)", level=2)
para(doc, "Antarmuka Landing Page menampilkan hero banner aplikasi IsyaratLive dengan penjelasan arsitektur 2-lapis AI (Computer Vision + LLM), informasi degradasi anggun offline, fitur panggilan WebRTC P2P, serta tombol navigasi utama untuk masuk ke Room Lokal, Room Remote, Kamus 32 Label, dan Uji Akurasi Model.")
add_image_with_caption(doc, LANDING_IMG, "Gambar 1. Antarmuka Landing Page IsyaratLive — Beranda utama aplikasi dengan visual light mode")

doc.add_paragraph()
heading(doc, "2. Tampilan Mode 2: Teks -> Isyarat (Pemutar Video Kalimat Isyarat)", level=2)
para(doc, "Mode 2 memungkinkan pengguna menerjemahkan teks atau ucapan menjadi rangkaian video peragaan isyarat BISINDO. Antarmuka dilengkapi dengan input teks, tombol speech-to-text, contoh kalimat cepat, pemutar video dengan kontrol kecepatan dan loop, serta pita alur kata kalimat (timeline).")
add_image_with_caption(doc, MODE2_IMG, "Gambar 2. Tampilan Mode 2: Teks -> Isyarat — pemutar video peragaan isyarat dan alur kata kalimat")

doc.add_paragraph()
heading(doc, "3. Tampilan Kamus Kosakata BISINDO (32 Label Terdaftar)", level=2)
para(doc, "Modal Kamus Kosakata menampilkan grid 32 kosakata BISINDO yang didukung oleh model AI. Setiap kartu kata memuat ID label, kategori, nama kata Bahasa Indonesia, terjemahan Bahasa Inggris, deskripsi, serta tombol 'Tonton Video' dan 'Tes' untuk menguji kata.")
add_image_with_caption(doc, DICT_IMG, "Gambar 3. Modal Kamus Kosakata BISINDO — grid 32 label terdaftar dengan filter kategori dan pencarian")

doc.add_paragraph()
heading(doc, "4. Tampilan Panel Uji Akurasi Model", level=2)
para(doc, "Panel interaktif untuk menguji performa model TensorFlow.js (Model v1 & v2) secara langsung di browser terhadap 32 video dictionary sebagai ground-truth sanity check.")
add_image_with_caption(doc, ACCURACY_IMG, "Gambar 4. Panel Uji Akurasi Model — evaluasi sanity check model TFJS langsung di browser")

# ============================================================
# i) DOKUMENTASI CARA PENGGUNAAN
# ============================================================
doc.add_page_break()
heading(doc, "i) Dokumentasi Cara Penggunaan Perangkat Lunak", level=1)

heading(doc, "Prasyarat", level=2)

add_styled_table(doc,
    ["Kebutuhan", "Detail"],
    [
        ["Perangkat", "Laptop, PC dengan webcam, atau smartphone"],
        ["Browser", "Google Chrome 90+, Mozilla Firefox 90+, atau Microsoft Edge 90+"],
        ["Koneksi Internet", "Diperlukan untuk fitur penyusunan kalimat natural (LLM). Tanpa internet, sistem berfungsi dalam mode degradasi anggun (gloss mentah secara lokal)"],
        ["Kamera", "Kamera webcam bawaan atau eksternal. Pastikan tidak diblokir oleh browser"],
        ["HTTPS", "Aplikasi diakses melalui HTTPS (akses kamera browser mensyaratkan HTTPS)"],
    ],
    col_widths=[4, 13]
)

doc.add_paragraph()
heading(doc, "Cara Menggunakan Aplikasi", level=2)

steps_usage = [
    "Buka aplikasi IsyaratLive melalui URL browser.",
    "Pilih mode pengujian melalui navigasi header: Beranda, Room Lokal, Room Remote, 32 Kata, atau Uji Akurasi.",
    "Pada Mode 2 (Teks -> Isyarat): Ketik kalimat atau pilih contoh kalimat (misal: 'Terima kasih'), lalu klik 'Tampilkan Video Kalimat' untuk memutar rangkaian video peragaan isyarat.",
    "Pada Kamus Kosakata: Klik '32 Kata' di header untuk menjelajahi grid 32 label BISINDO berdasarkan kategori atau kolom pencarian.",
    "Pada Panel Uji Akurasi: Klik 'Uji Akurasi' untuk menjalankan sanity check model TFJS langsung di browser.",
]
for i, step in enumerate(steps_usage, 1):
    bullet(doc, f"{i}. {step}")

# ============================================================
# KELEBIHAN DIBANDING YANG SERUPA
# ============================================================
doc.add_page_break()
heading(doc, "Kelebihan IsyaratLive Dibandingkan Perangkat Lunak Serupa", level=1)

add_styled_table(doc,
    ["Aspek", "IsyaratLive", "Solusi Sejenis yang Ada"],
    [
        ["Bahasa isyarat", "BISINDO (bahasa alami komunitas Tuli)", "Mayoritas SIBI (konstruksi artifisial)"],
        ["Level deteksi", "Rangkaian kata dinamis (32 kosakata level-kata)", "Alfabet/angka statis atau kata tunggal"],
        ["Pemahaman bahasa", "Dua lapis AI: CV + LLM untuk kalimat natural", "Hanya CV, output kata lepas tanpa struktur"],
        ["Dua arah", "Isyarat->Teks DAN Teks->Isyarat", "Umumnya satu arah saja"],
        ["Platform", "Web browser (tanpa instalasi, cukup buka URL)", "Sering memerlukan instalasi aplikasi"],
        ["Offline capability", "Mode degradasi anggun (gloss mentah tetap ditampilkan)", "Umumnya tidak berfungsi tanpa internet"],
        ["Testing & Quality", "Test suite Vitest (31 FE & 10 BE test passing)", "Umumnya tanpa otomatisasi testing"],
        ["Panggilan Remote", "Panggilan video WebRTC P2P 1-lawan-1 terintegrasi", "Tidak ada fitur video call terintegrasi"],
    ],
    col_widths=[4, 6.5, 6.5]
)

# ============================================================
# DAFTAR KOMPONEN DAN LISENSI
# ============================================================
doc.add_paragraph()
heading(doc, "Daftar Komponen dan Lisensi", level=1)

add_styled_table(doc,
    ["Komponen / Library", "Versi", "Lisensi", "Keterangan"],
    [
        ["React", "19.2.8", "MIT", "Library UI frontend"],
        ["Vite", "8.2.0", "MIT", "Build tool & dev server"],
        ["Tailwind CSS", "4.3.3", "MIT", "CSS framework"],
        ["TypeScript", "6.0.2 / 7.0.2", "Apache-2.0", "Bahasa pemrograman"],
        ["@mediapipe/tasks-vision", "1.0.1", "Apache-2.0", "Deteksi landmark tangan (Google)"],
        ["@tensorflow/tfjs", "4.22.0", "Apache-2.0", "Inferensi model ML di browser"],
        ["Express", "5.2.1", "MIT", "Framework backend Node.js"],
        ["mysql2", "3.23.2", "MIT", "Driver MySQL untuk Node.js"],
        ["dotenv", "17.4.2", "BSD-2-Clause", "Environment variable loader"],
        ["cors", "2.8.6", "MIT", "Middleware CORS"],
        ["Socket.io", "4.8.3", "MIT", "Signaling Room Remote (WebRTC)"],
        ["Vitest", "4.1.10", "MIT", "Framework unit testing"],
        ["MediaPipe (Python)", "0.10.14+", "Apache-2.0", "Ekstraksi landmark saat training"],
        ["TensorFlow (Python)", "2.15.1+", "Apache-2.0", "Training model"],
        ["TensorFlow.js Converter", "4.20.0+", "Apache-2.0", "Konversi model ke format web"],
        ["scikit-learn", "1.6.1", "BSD-3-Clause", "Utilitas ML (evaluasi, split)"],
        ["OpenCV (Python)", "4.11.0", "Apache-2.0", "Pemrosesan video saat training"],
        ["NumPy", "2.2.3", "BSD-3-Clause", "Komputasi numerik"],
        ["Pandas", "2.2.3", "BSD-3-Clause", "Manipulasi data tabular"],
        ["Dataset WL-BISINDO", "2025", "CC BY-NC 4.0", "Dataset pelatihan (non-komersial, wajib sitasi)"],
    ],
    col_widths=[4.5, 3, 3, 6.5]
)

# ============================================================
# REFERENSI
# ============================================================
doc.add_paragraph()
heading(doc, "Referensi dan Sitasi", level=1)

heading(doc, "Dataset WL-BISINDO", level=2)
para(doc, "Kindy, G.O., Leonali, G., & Lucky, H. (2025). WL-BISINDO: A Word-Level Bahasa Isyarat Indonesia Dataset. Procedia Computer Science. DOI: 10.1016/j.procs.2025.08.277. Elsevier.", size=10)
para(doc, "Sumber dataset: https://www.kaggle.com/datasets/glennleonali/wl-bisindo", size=10)

heading(doc, "Teknologi Utama", level=2)
bullet(doc, "Google MediaPipe — https://developers.google.com/mediapipe")
bullet(doc, "TensorFlow.js — https://www.tensorflow.org/js")
bullet(doc, "Web Speech API — https://developer.mozilla.org/en-US/docs/Web/API/Web_Speech_API")
bullet(doc, "React — https://react.dev")
bullet(doc, "Vite — https://vite.dev")
bullet(doc, "Express.js — https://expressjs.com")

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— Dokumen proposal ini disusun untuk kompetisi GEMASTIK XIX 2026, Divisi VIII: Pengembangan Perangkat Lunak —")
run.font.size = Pt(9)
run.italic = True
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ============================================================
# SAVE TO BOTH PATHS
# ============================================================
doc.save(OUTPUT_PATH_1)
print(f"[OK] Proposal DOCX berhasil dibuat: {OUTPUT_PATH_1}")
print(f"     Ukuran: {os.path.getsize(OUTPUT_PATH_1) / 1024:.1f} KB")

doc.save(OUTPUT_PATH_2)
print(f"[OK] Proposal DOCX berhasil dibuat: {OUTPUT_PATH_2}")
print(f"     Ukuran: {os.path.getsize(OUTPUT_PATH_2) / 1024:.1f} KB")
